#!/usr/bin/env python3
"""Build the editable MATH 346 project report templates.

The documents deliberately use a restrained Google Docs-compatible style:
Arial typography, black hierarchy, simple tables only for repeated fields, and
short muted prompts that students replace with their own work.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_SCRIPTS = Path(
    "/Users/laurentnajman/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.715.12143/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


PAGE_WIDTH_DXA = 9360
BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
BORDER_HEX = "DADCE0"
FONT_NAME = "Arial"


def set_run_font(run, *, size=None, bold=None, italic=None, color=BLACK):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, *, color=BORDER_HEX, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_numbering(document, *, num_id, abstract_id, bullet):
    numbering = document.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "●" if bullet else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT_NAME)
        r_fonts.set(qn("w:hAnsi"), FONT_NAME)
        r_pr.append(r_fonts)
        level.append(r_pr)

    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def apply_num_id(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (20, BLACK, 20, 6),
        "Heading 2": (16, BLACK, 18, 6),
        "Heading 3": (14, RGBColor(0x43, 0x43, 0x43), 16, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    prompt = document.styles.add_style("Template Prompt", 1)
    prompt.base_style = normal
    prompt.font.name = FONT_NAME
    prompt._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    prompt._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    prompt.font.size = Pt(10.5)
    prompt.font.italic = True
    prompt.font.color.rgb = MUTED
    prompt.paragraph_format.space_after = Pt(10)
    prompt.paragraph_format.line_spacing = 1.15

    caption = document.styles["Caption"]
    caption.font.name = FONT_NAME
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.15

    add_numbering(document, num_id=91, abstract_id=91, bullet=True)
    add_numbering(document, num_id=92, abstract_id=92, bullet=False)

    document.core_properties.title = ""
    document.core_properties.subject = "MATH 346 project report template"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = "Editable report template for students"


def add_title(document, title, subtitle):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    set_run_font(run, size=26, bold=False)

    subtitle_p = document.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(16)
    subtitle_p.paragraph_format.keep_with_next = True
    run = subtitle_p.add_run(subtitle)
    set_run_font(run, size=13, color=MUTED)


def add_table(document, headers, rows, widths_dxa, *, font_size=10):
    table = document.add_table(rows=1, cols=len(headers))
    table.allow_autofit = False
    header = table.rows[0]
    mark_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(text)
        set_run_font(run, size=font_size, bold=True)

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(value)
            is_prompt = value.startswith("[") and value.endswith("]")
            set_run_font(
                run,
                size=font_size,
                italic=is_prompt,
                color=MUTED if is_prompt else BLACK,
            )

    # Keep short data-entry tables intact when possible. The table rows remain
    # expandable, but Word/LibreOffice should not leave a final row alone on
    # the next page.
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        if row_index < len(table.rows) - 1:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True

    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=0,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    set_table_borders(table)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_metadata(document, project_name):
    add_table(
        document,
        ["Field", "Complete this"],
        [
            ("Project", project_name),
            ("Group identifier", "[group-__]"),
            ("Student names", "[Name 1; Name 2; Name 3]"),
            ("Submission date", "[Day Month Year]"),
        ],
        [2160, 7200],
        font_size=10.5,
    )


def add_paragraph(document, text="", *, bold_label=None, keep=False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = keep
    if bold_label:
        run = paragraph.add_run(bold_label)
        set_run_font(run, bold=True)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_prompt(document, text):
    return document.add_paragraph(f"[Replace this prompt: {text}]", style="Template Prompt")


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    apply_num_id(paragraph, 91)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_numbered(document, text):
    paragraph = document.add_paragraph()
    apply_num_id(paragraph, 92)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_checkbox(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"[ ]  {text}")
    set_run_font(run, size=10.5)
    return paragraph


def add_label_response(document, label, prompt):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(label)
    set_run_font(run, bold=True)
    add_prompt(document, prompt)


def add_figure_slot(document, number, purpose, filename):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"Figure {number}. {purpose}")
    set_run_font(run, bold=True)
    prompt = document.add_paragraph(style="Template Prompt")
    prompt.paragraph_format.space_after = Pt(4)
    run = prompt.add_run(
        f"[Insert {filename} here. Resize it so labels remain readable. Delete this instruction.]"
    )
    set_run_font(run, size=10.5, italic=True, color=MUTED)
    caption = document.add_paragraph(style="Caption")
    caption.add_run(
        f"Figure {number}. [Write a one-sentence caption stating what the figure shows and the main point.]"
    )
    add_prompt(document, "Interpret the figure in 2-4 sentences; do not merely repeat the axes.")


def add_page_break(document):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_common_front_matter(document, *, project_name, report_filename, source_files, output_files):
    add_metadata(document, project_name)
    add_paragraph(
        document,
        "This template is designed for Microsoft Word or Google Docs. Replace every gray italic prompt with your own writing, insert the figures generated by your code, and delete all remaining instructions before export.",
    )

    document.add_heading("How to submit", level=1)
    for text in (
        "Replace 01 in every example filename with the two-digit number assigned to your group.",
        "Keep the report focused on results and interpretation; do not paste long blocks of source code into it.",
        "Check that every number in the report agrees with the files generated by your script.",
        f"Export this document as {report_filename} in PDF format.",
        "Submit the PDF, source files, and generated CSV outputs using the exact filenames shown below.",
    ):
        add_numbered(document, text)

    add_table(
        document,
        ["File type", "Required filename"],
        [(kind, filename) for kind, filename in source_files + output_files],
        [2300, 7060],
        font_size=10,
    )

    add_paragraph(
        document,
        "Possible verification. Any group member may be selected briefly, either randomly or because the submitted work requires clarification, to explain code, a figure, a numerical result, or a contribution. Selection is not itself an accusation of misconduct.",
        bold_label="",
    )


def add_ai_and_contributions(document, project_label, *, checklist_new_page=False):
    document.add_heading("AI-use appendix", level=1)
    add_paragraph(
        document,
        "Record material generative-AI assistance. Describe what your group checked, changed, or rejected. If no generative AI was used, write that explicitly in the first row.",
    )
    add_table(
        document,
        ["Tool", "Task or prompt", "Output used?", "Changed or rejected", "Student check"],
        [
            ("[Tool or none]", "[Concise description]", "[Yes / partly / no]", "[What changed]", "[How verified]"),
            ("[Optional]", "[Optional]", "[Optional]", "[Optional]", "[Optional]"),
            ("[Optional]", "[Optional]", "[Optional]", "[Optional]", "[Optional]"),
        ],
        [1100, 2350, 1150, 2650, 2110],
        font_size=8.5,
    )

    document.add_heading("Group contribution statement", level=1)
    add_paragraph(
        document,
        "Each member should state a principal contribution and identify another member who checked it. By submitting, all members confirm that they reviewed the final files and accept responsibility for understanding the work.",
    )
    add_table(
        document,
        ["Student", "Principal contribution", "Work checked by"],
        [
            ("[Name 1]", "[Contribution]", "[Name]"),
            ("[Name 2]", "[Contribution]", "[Name]"),
            ("[Name 3]", "[Contribution]", "[Name]"),
        ],
        [1900, 5300, 2160],
        font_size=9.5,
    )

    if checklist_new_page:
        add_page_break(document)
    document.add_heading("Final submission check", level=1)
    for text in (
        "All gray prompts and unused instructions have been deleted.",
        "The PDF names the correct group and project.",
        "The complete script runs from beginning to end in the project folder.",
        "Figures in the PDF were generated by the submitted code.",
        "Reported values agree with the generated CSV outputs.",
        "All required filenames are exact.",
        "Material AI assistance is documented and checked.",
        f"Every group member has reviewed the complete {project_label} submission.",
    ):
        add_checkbox(document, text)


def build_project_1(path):
    document = Document()
    configure_document(document)
    add_title(
        document,
        "Project 1 report",
        "MATLAB numerical investigation of synthetic building-energy data",
    )
    add_common_front_matter(
        document,
        project_name="Project 1: MATLAB Numerical Investigation",
        report_filename="group-01-project-1.pdf",
        source_files=[
            ("MATLAB main script", "group_01_project1.m"),
            ("MATLAB function", "model_rmse.m"),
        ],
        output_files=[("Generated results", "group-01-project1-results.csv")],
    )

    add_page_break(document)
    document.add_heading("1. Problem and assigned dataset", level=1)
    add_prompt(
        document,
        "In one short paragraph, explain the facilities-management question, state that the data are synthetic, and identify your assigned group/building."
    )
    add_label_response(
        document,
        "Assigned prediction temperature: ",
        "Report the temperature from the assignment CSV, with units, and confirm that it lies inside the usable temperature range.",
    )

    document.add_heading("2. Data checks", level=1)
    add_prompt(
        document,
        "Describe how the CSV was imported and checked. Report the raw and usable row counts, the missing value, variable ranges, and any unusual observations you found. Explain one logical-indexing check."
    )
    add_table(
        document,
        ["Check", "Result from MATLAB"],
        [
            ("Raw rows", "[Value]"),
            ("Usable rows", "[Value]"),
            ("Mean outdoor temperature (°C)", "[Value]"),
            ("Mean heating energy (kWh)", "[Value]"),
            ("Logical-indexing condition", "[Condition and row count]"),
        ],
        [4000, 5360],
        font_size=10,
    )

    document.add_heading("3. Function check", level=1)
    add_prompt(
        document,
        "State the hand-checkable model_rmse test, its expected value, and how the successful assertion supports confidence in the function."
    )

    add_page_break(document)
    document.add_heading("4. Figures", level=1)
    add_figure_slot(document, 1, "Heating energy over time or observation order", "group-01-project1-series.png")
    add_figure_slot(document, 2, "Heating energy against outdoor temperature with fitted curves", "group-01-project1-fit.png")
    add_figure_slot(document, 3, "Residuals from the quadratic fit", "group-01-project1-residuals.png")

    add_page_break(document)
    document.add_heading("5. Linear and quadratic fits", level=1)
    add_paragraph(
        document,
        "Use mathematical coefficient order in this report: intercept, temperature, then temperature squared. MATLAB polyfit returns the reverse storage order, so check the mapping carefully.",
    )
    add_table(
        document,
        ["Model", "Intercept", "Temp.", "Temp.²", "RMSE (kWh)", "R²"],
        [
            ("Linear", "[b0]", "[b1]", "not used", "[Value]", "[Value]"),
            ("Quadratic", "[b0]", "[b1]", "[b2]", "[Value]", "[Value]"),
        ],
        [1440, 1400, 1400, 1400, 1980, 1740],
        font_size=9,
    )
    add_label_response(
        document,
        "Model comparison: ",
        "Compare RMSE, R-squared, curve shape, and residual behavior. Explain whether the extra quadratic term materially improves the fit.",
    )
    add_label_response(
        document,
        "Unusual observations: ",
        "Identify observations with large residuals and distinguish an unusual point from a data-entry error.",
    )
    add_label_response(
        document,
        "Preferred descriptive model: ",
        "State whether you prefer the linear or quadratic fit for description and justify the choice. The required contract prediction below still uses the quadratic fit.",
    )

    document.add_heading("6. Assigned prediction", level=1)
    add_table(
        document,
        ["Quantity", "Required value"],
        [
            ("Assigned temperature (°C)", "[Value]"),
            ("Quadratic prediction (kWh)", "[Value]"),
            ("Observed usable temperature range (°C)", "[Minimum to maximum]"),
        ],
        [5100, 4260],
        font_size=10,
    )
    add_prompt(
        document,
        "Interpret the prediction in context. Explain why it is interpolation and why using the fitted curve outside the observed range could be unsafe."
    )

    document.add_heading("7. Sensitivity check", level=1)
    add_table(
        document,
        ["Quantity", "Required value"],
        [
            ("Removed observation_id", "[Identifier]"),
            ("Original quadratic prediction (kWh)", "[Value]"),
            ("Refitted quadratic prediction (kWh)", "[Value]"),
            ("Absolute or percentage change", "[Value and units]"),
        ],
        [5100, 4260],
        font_size=10,
    )
    add_prompt(
        document,
        "Explain how the removed row was selected and whether the substantive conclusion survives this fixed sensitivity check."
    )

    document.add_heading("8. Conclusion and limitations", level=1)
    add_prompt(
        document,
        "Write a concise conclusion answering the project question. Include the main numerical result, one limitation of the synthetic data or fitted relationship, and one appropriate next check. Avoid causal claims."
    )

    add_page_break(document)
    add_ai_and_contributions(document, "Project 1", checklist_new_page=True)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def build_project_2(path):
    document = Document()
    configure_document(document)
    add_title(
        document,
        "Project 2 report",
        "R modeling and validation with synthetic building-energy data",
    )
    add_common_front_matter(
        document,
        project_name="Project 2: R Modeling and Validation",
        report_filename="group-01-project-2.pdf",
        source_files=[("R source file", "group-01-project2.R")],
        output_files=[
            ("Generated results", "group-01-project2-results.csv"),
            ("Generated predictions", "group-01-validation-predictions.csv"),
        ],
    )

    add_page_break(document)
    document.add_heading("1. Question and assigned data", level=1)
    add_prompt(
        document,
        "In one short paragraph, refine the facilities-management question, state that the data are synthetic, and identify your assigned group/building. Name the response and candidate predictors."
    )

    document.add_heading("2. Data checks and cleaning", level=1)
    add_prompt(
        document,
        "Describe how the training CSV was imported and checked. Report data types, missingness, the repeated observation_id, ranges, and categories. Explain why cleaning is performed through code rather than by editing the CSV."
    )
    add_table(
        document,
        ["Cleaning checkpoint", "Row count"],
        [
            ("Raw training rows", "[Value]"),
            ("Unique observation_id rows", "[Value]"),
            ("Complete usable training rows", "[Value]"),
        ],
        [5900, 3460],
        font_size=10,
    )

    document.add_heading("3. Transformations and grouped summaries", level=1)
    add_prompt(
        document,
        "Describe the transformations used and report at least one grouped summary that informs the analysis. Include a row-count check after an important transformation."
    )

    add_page_break(document)
    document.add_heading("4. Exploratory figures", level=1)
    add_figure_slot(document, 1, "Distribution of an important variable", "group-01-project2-distribution.png")
    add_figure_slot(document, 2, "Relationships relevant to the proposed model", "group-01-project2-relationships.png")

    add_page_break(document)
    document.add_heading("5. Frozen model plan", level=1)
    add_paragraph(
        document,
        "Complete this section before the validation file is released. Submit or record this page in the form announced by the instructor. Do not inspect validation outcomes before the plan is frozen.",
    )
    add_label_response(document, "Numeric response: ", "Name the response variable and units.")
    add_label_response(document, "Selected predictors: ", "List the variables and any approved transformations or interactions.")
    add_label_response(document, "Exact R formula: ", "Paste the one-line formula used by model_formula in the submitted script.")
    add_label_response(document, "Principal quantity: ", "Name one exact R coefficient and state the contextual interpretation you intend to make.")
    add_label_response(document, "Baseline prediction: ", "State the simple baseline and why it is a meaningful comparison.")
    add_label_response(document, "Validation metric: ", "Identify the primary metric and explain whether lower or higher is better.")
    add_label_response(document, "Expected failure condition: ", "Describe one situation in which the model may predict poorly.")
    add_label_response(document, "Plan frozen on: ", "Enter the date and time of the Blackboard record or submission.")
    add_checkbox(document, "The validation file had not been opened or analyzed when this plan was frozen.")

    add_page_break(document)
    document.add_heading("6. Training fit and diagnostics", level=1)
    add_table(
        document,
        ["Training quantity", "Value"],
        [
            ("Frozen formula", "[Exact formula]"),
            ("Primary coefficient name", "[Exact R name]"),
            ("Primary coefficient estimate", "[Value]"),
            ("Training RMSE (kWh)", "[Value]"),
            ("Training MAE (kWh)", "[Value]"),
            ("Training R²", "[Value]"),
        ],
        [5100, 4260],
        font_size=9.5,
    )
    add_label_response(
        document,
        "Coefficient interpretation: ",
        "Interpret the selected coefficient in context, including units and any conditions required by transformations, interactions, or reference categories.",
    )
    add_figure_slot(document, 3, "Residual and influence diagnostics", "group-01-project2-residuals.png")
    add_label_response(
        document,
        "Diagnostic assessment: ",
        "Identify important residual patterns or influential observations and state which modeling assumptions matter for your conclusion.",
    )

    add_page_break(document)
    document.add_heading("7. Validation results", level=1)
    add_table(
        document,
        ["Evaluation", "RMSE (kWh)", "MAE (kWh)", "R²"],
        [
            ("Mean-response baseline on validation", "[Value]", "not required", "not required"),
            ("Frozen model on training", "[Value]", "[Value]", "[Value]"),
            ("Frozen model on validation", "[Value]", "[Value]", "not required"),
        ],
        [3540, 2100, 1980, 1740],
        font_size=8.8,
    )
    add_prompt(
        document,
        "Compare baseline, training, and validation performance. Quantify the improvement over the baseline and the change from training to validation. Explain whether the frozen model generalizes adequately."
    )
    add_label_response(
        document,
        "Prediction alignment check: ",
        "Confirm that the generated prediction CSV has one row per validation observation_id in validation-file order."
    )
    add_label_response(
        document,
        "Where predictions fail: ",
        "Describe the conditions or observations with the largest validation errors, without redesigning the frozen formula around them."
    )

    document.add_heading("8. Sensitivity analysis", level=1)
    add_table(
        document,
        ["Quantity", "Required value"],
        [
            ("Removed training observation_id", "[Identifier]"),
            ("Original validation RMSE (kWh)", "[Value]"),
            ("Sensitivity validation RMSE (kWh)", "[Value]"),
            ("Absolute or percentage change", "[Value and units]"),
        ],
        [5100, 4260],
        font_size=10,
    )
    add_prompt(
        document,
        "Explain how the row was selected using the largest absolute standardized residual, confirm that the same formula was refit, and state whether the substantive conclusion survives."
    )

    document.add_heading("9. Conclusion and limitations", level=1)
    add_prompt(
        document,
        "Answer the refined question using the validation evidence. Include one limitation, one appropriate next check, and no unsupported causal claims."
    )

    add_page_break(document)
    add_ai_and_contributions(document, "Project 2", checklist_new_page=True)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    build_project_1(args.output_dir / "project1-report-template.docx")
    build_project_2(args.output_dir / "project2-report-template.docx")
    print(f"Built report templates in {args.output_dir}")


if __name__ == "__main__":
    main()
